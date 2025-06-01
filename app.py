from flask import Flask, render_template, request, redirect, url_for, send_from_directory, flash
from werkzeug.utils import secure_filename
from datetime import datetime
from docx import Document
import re
import os


app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['SECRET_KEY'] = 'yqfm_yurist_2025'
app.config['ALLOWED_EXTENSIONS'] = {'docx'}

# labels for each template
list_labels = {
    "temp1": ['name1', 'date1'],
    "temp2": ['name2', 'date2'],
    "temp3": ['name3', 'date3'],
}

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Allowed file type check
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

# Clean and timestamp filename
def clean_filename(filename):
    name, ext = os.path.splitext(secure_filename(filename))
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{name}_{timestamp}{ext}"


def replace_text_in_docx(input_path, replacements, output_path):
    doc = Document(input_path)
    for paragraph in doc.paragraphs:
        full_text = ''.join(run.text for run in paragraph.runs)
        for old, new in replacements.items():
            if old in full_text:
                full_text = full_text.replace(old, new)
                for i in range(len(paragraph.runs) - 1, -1, -1):
                    p_run = paragraph.runs[i]
                    p_run.text = ''
                paragraph.runs[0].text = full_text

    doc.save(output_path)


def extract_text_and_placeholders(file_path):
    doc = Document(file_path)
    text = "\n".join([p.text for p in doc.paragraphs])

    placeholders = re.findall(r"\{(.*?)\}", text)
    return text, list(set(placeholders))

# Homepage
@app.route('/', methods=['GET', 'POST'])
def index():
    return render_template('index.html')

# Document processing route
@app.route('/document/<name>', methods=['GET', 'POST'])
def download(name):
    name_ = str(name).split(".")[0]
    template_path = os.path.abspath(f'uploads/{name}.docx')
    text, placeholders = extract_text_and_placeholders(template_path)
    text_temp = {"text": text}

    if request.method == 'POST':
        replacements = {label: request.form[label] for label in list_labels.get(name_, [])}
        modified_filename = f"modified_{name}.docx"
        modified_path = os.path.join(app.config['UPLOAD_FOLDER'], modified_filename)
        replace_text_in_docx(template_path, replacements, modified_path)
        return send_from_directory(app.config['UPLOAD_FOLDER'], modified_filename, as_attachment=True)

    # GET request: show the form
    return render_template('document.html', labels=list_labels.get(name_, []), file_name=name, template=text_temp)

# @app.errorhandler(413)
# def request_entity_too_large(error):
#     flash('File too large (max 16MB)', 'error')
#     return redirect(url_for('index'))

# Start server
if __name__ == '__main__':
    app.run(debug=True)
